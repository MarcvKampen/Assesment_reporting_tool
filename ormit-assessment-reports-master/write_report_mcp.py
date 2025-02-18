import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import ast
import json
import re
import logging

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    filename="report_generation_mcp.log",  # Different log file
    filemode="a"
)

DETAILS_TABLE_INDEX = 0
COGCAP_TABLE_INDEX = 1
CONCLUSION_TABLE_INDEX = 2
FIRST_ICONS_TABLE = 4
NUM_ICONS_TABLES = 5
ITEMS_PER_ICON_TABLE = 4
LANGUAGE_SKILLS_TABLE_INDEX = 13 # This is no longer used for table index

def resource_path(relative_path):
    """Get absolute path to resource, works for dev and PyInstaller."""
    try:
        base_path = sys._MEIPASS
    except AttributeError:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

def _safe_get_table(doc, table_index, default=None):
    """Safely retrieves a table."""
    try:
        return doc.tables[table_index]
    except IndexError:
        logging.warning(f"Table {table_index} not found.")
        return default

def _safe_get_cell(table, row_index, col_index, default=None):
    """Safely retrieves a cell."""
    try:
        return table.cell(row_index, col_index)
    except IndexError:
        logging.warning(f"Cell ({row_index}, {col_index}) not found.")
        return default

def _safe_set_text(cell, text):
    """Safely sets cell text, clearing existing content."""
    if cell:
        for p in cell.paragraphs:
            p = p._element
            p.getparent().remove(p)
        paragraph = cell.add_paragraph()
        run = paragraph.add_run(str(text))
        run.font.name = 'Montserrat Light'
        run.font.size = Pt(10)

def _safe_add_paragraph(cell, text):
    """Safely add paragraphs"""
    if cell:
        paragraph = cell.add_paragraph(text)
        run = paragraph.runs[0]
        run.font.name = 'Montserrat Light'
        run.font.size = Pt(10)

        r = run._element
        rPr = r.rPr
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.append(rPr)

        rFonts = OxmlElement('w:rFonts')
        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', 'Montserrat Light')
        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi', 'Montserrat Light')
        rPr.append(rFonts)

def _safe_literal_eval(s, default=None):
    """Safely evaluates a string as a Python literal."""
    try:
        return ast.literal_eval(s)
    except (SyntaxError, ValueError) as e:
        logging.error(f"Error evaluating string: {s} - {e}")
        return default

def find_and_replace_placeholder(doc, placeholder, replacement_text, font_name='Montserrat Light', font_size=10):
    """Finds and replaces ALL occurrences of a placeholder in the document (paragraphs and tables)."""
    placeholder_found = False

    # 1. Search Paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs: # Iterate through runs in each paragraph
            if placeholder in run.text:
                placeholder_found = True
                run.text = run.text.replace(placeholder, str(replacement_text))
                run.font.name = font_name
                run.font.size = Pt(font_size)

    # 2. Search Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    placeholder_found = True
                    _safe_set_text(cell, replacement_text) # Use safe set (clears and sets font)

    if not placeholder_found:
        logging.warning(f"Placeholder '{placeholder}' not found.")

# --- Main Functions ---
def clean(text):
    """Cleans input text."""
    return re.sub(r'[\【】`]|(```python)|(\*\*)', '', str(text)).strip() if isinstance(text, str) else text

def strip_extra_quotes(input_string):
    """Removes leading/trailing double quotes."""
    if isinstance(input_string, str) and input_string.startswith('"') and input_string.endswith('"'):
        return input_string[1:-1]
    return input_string

def clean_up(loc_dic):
    """Loads and cleans JSON data."""
    try:
        with open(loc_dic, 'r', encoding='utf-8') as f:
            loaded_data = json.load(f)
        return {key: clean(value) for key, value in loaded_data.items()}
    except (FileNotFoundError, json.JSONDecodeError) as e:
        logging.error(f"Error loading/cleaning JSON: {e}")
        return {}

def update_document(output_dic, name, assessor, gender, program):
    """Updates the Word document (MCP version)."""
    try:
        doc = Document(resource_path('resources/template.docx'))  # MCP Template
    except Exception as e:
        logging.error(f"Failed to open template: {e}")
        return None

    # --- Static Content ---
    add_content_detailstable(doc, [name, "", program, "", ""])
    replace_and_format_header_text(doc, name)
    find_and_replace_placeholder(doc, '***', name.split()[0], 'Montserrat Light') # Using find_and_replace_placeholder
    find_and_replace_placeholder(doc, 'ASSESSOR', assessor.upper(), 'Montserrat Light') # Using find_and_replace_placeholder

    # --- Dynamic Content ---
    dynamic_prompts = [
        'prompt2_firstimpr', 'prompt3_personality',
        'prompt4_cogcap_remarks',         'prompt4_cogcap_remarks', 'prompt6a_conqual',
        'prompt6b_conimprov', 'prompt9_interests'
    ]
    for prompt_key in dynamic_prompts:
      replacement = output_dic.get(prompt_key, "")
      if prompt_key in ['prompt2_firstimpr','prompt3_personality','prompt4_cogcap_remarks']:
          replacement = replacePiet(replacement,name, gender) # Apply replacePiet
      find_and_replace_placeholder(doc, f"{{{prompt_key}}}", replacement)


    # ---  Content that remains in specific locations (no longer tables for language skills) ---
    add_content_cogcaptable(doc, output_dic.get('prompt4_cogcap_scores', "[]"))
    language_skills(doc, output_dic.get('prompt5_language', "[]"), doc) # Pass doc object

    # Profile review (icons)
    qual_scores_str = output_dic.get('prompt7_qualscore', "[]")
    qual_scores = _safe_literal_eval(qual_scores_str, [])

    if isinstance(qual_scores, list):
            add_icons2(doc, qual_scores)  # Use the MCP version
    else:
        logging.warning("Invalid qual_scores data.")

    # --- Conclusion Table ---
    conclusion(doc, 0, output_dic.get('prompt6a_conqual', "[]"))  # Strengths to column 0
    conclusion(doc, 1, output_dic.get('prompt6b_conimprov', "[]"))  # Improvements to column 1


    # --- Save Document ---
    current_time = datetime.now()
    formatted_time = current_time.strftime("%m%d%H%M")
    updated_doc_path = f"Assessment Report - {name} - {formatted_time}.docx"
    try:
        doc.save(updated_doc_path)
        return updated_doc_path
    except Exception as e:
        logging.error(f"Failed to save document: {e}")
        return None

def format_datatools_output(datatools_json_string):
    """Formats data tools output (not used in MCP, kept for consistency)."""
    try:
        return "\n".join(f"- {tool}: {level}" for tool, level in ast.literal_eval(datatools_json_string).items())
    except (ValueError, SyntaxError):
        return "Could not parse data tools information."

def format_interests_output(interests_json_string):
    """Formats interests output (not directly used in MCP, kept for consistency)."""
    try:
        return "\n".join(f"- {interest}" for interest in ast.literal_eval(interests_json_string))
    except (ValueError, SyntaxError):
        return "Could not parse interests information."

def replacePiet(text, name, gender):
    """Replaces 'Piet' and handles pronouns."""
    if not isinstance(text, str):
        return ""

    text = text.replace("Piet", name.split()[0])
    text = re.sub(r'\bthe trainee\b', name.split()[0], text, flags=re.IGNORECASE)

    if gender == 'M':
        replacements = {"She": "He", "she": "he", "Her": "Him", "her": "him",
                        "Hers": "His", "hers": "his", "Herself": "Himself", "herself": "himself"}
    elif gender == 'F':
        replacements = {"He": "She", "he": "she", "Him": "Her", "him": "her",
                        "His": "Her", "his": "her", "Himself": "Herself", "himself": "herself"}
    else:
        return text

    for from_word, to_word in replacements.items():
        text = re.sub(r'\b' + re.escape(from_word) + r'\b', to_word, text)
    return text

def restructure_date(date_str):
    """Restructures date to DD-MM-YYYY."""
    date_str = date_str.replace('/', '-')
    try:
        datetime.strptime(date_str, '%d-%m-%Y')
        return date_str
    except ValueError:
        try:
            return datetime.strptime(date_str, '%Y-%m-%d').strftime('%d-%m-%Y')
        except ValueError:
            return ''
def set_font_properties(cell):
    """Sets font properties for a cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Montserrat Light'
            run.font.size = Pt(11)
            r = run._element
            rPr = r.rPr
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.append(rPr)
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Montserrat Light')
            rFonts.set(qn('w:hAnsi'), 'Montserrat Light')
            rPr.append(rFonts)

def set_font_properties2(para):
    """Sets font properties with tabs for language skills."""
    full_text = para.text
    para.clear()
    lines = full_text.splitlines()

    for line in lines:
        words = line.split()
        if words:
            for word in words[:-1]:
                run = para.add_run(word + ' ')
                run.font.name = 'Montserrat Light'
                run.font.size = Pt(10)
                run.bold = False
                r = run._element
                rPr = r.rPr or OxmlElement('w:rPr')
                r.append(rPr)
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), 'Montserrat Light')
                rFonts.set(qn('w:hAnsi'), 'Montserrat Light')
                rPr.append(rFonts)

            if words[0] == 'Dutch':
                para.add_run('\t\t')
            else:
                para.add_run('\t')

            last_word = words[-1]
            last_run = para.add_run(last_word)
            last_run.font.name = 'Montserrat Light'
            last_run.font.size = Pt(10)
            last_run.bold = True
            r = last_run._element
            rPr = r.rPr or OxmlElement('w:rPr')
            r.append(rPr)
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Montserrat Light')
            rFonts.set(qn('w:hAnsi'), 'Montserrat Light')
            rPr.append(rFonts)

def add_content_detailstable(doc, personal_details):
    table = _safe_get_table(doc, DETAILS_TABLE_INDEX)
    if not table:
        return

    if not isinstance(personal_details, list):
        logging.warning("personal_details is not a list.")
        return

    if len(personal_details) == 1 and all(isinstance(ele, str) for ele in personal_details):
        personal_details = personal_details[0].split(',')

    for row_index, row in enumerate(table.rows):
        if len(row.cells) > 1:
            first_cell_text = row.cells[0].text.strip()
            second_cell_text = row.cells[1].text.strip()

            if first_cell_text == "Name candidate" and second_cell_text == ":":
                cell = _safe_get_cell(table, row_index, 2)
                _safe_set_text(cell, personal_details[0] if len(personal_details) > 0 else '')

            if first_cell_text == "Date of birth" and second_cell_text == ":":
                cell = _safe_get_cell(table, row_index, 2)
                _safe_set_text(cell, restructure_date(personal_details[1]) if len(personal_details) > 1 else '')

            if first_cell_text == "Position" and second_cell_text == ":":
                cell = _safe_get_cell(table, row_index, 2)
                _safe_set_text(cell, personal_details[2] if len(personal_details) > 2 else '')

            if first_cell_text == "Assessment date" and second_cell_text == ":":
                cell = _safe_get_cell(table, row_index, 2)
                _safe_set_text(cell, restructure_date(personal_details[3]) if len(personal_details) > 3 else '')

            if first_cell_text == "Pool" and second_cell_text == ":":
                cell = _safe_get_cell(table, row_index, 2)
                _safe_set_text(cell, personal_details[4] if len(personal_details) > 4 else '')

def add_content_cogcaptable(doc, scores_str):
    """Adds cognitive capacity scores."""
    table = _safe_get_table(doc, COGCAP_TABLE_INDEX)
    if not table:
        return

    scores = _safe_literal_eval(scores_str, [])
    if not isinstance(scores, list) or len(scores) != 6:
        logging.warning("Invalid scores data. Expected a list of 6 numbers.")
        return

    for i in range(6):
        cell = _safe_get_cell(table, 1, i + 1)  # Row 1 (second row)
        if cell:
            if i == 0:
                _safe_set_text(cell, scores[i])
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.bold = True
                run.underline = True
                paragraph.alignment = 1
            else:
                _safe_set_text(cell, scores[i])
                paragraph = cell.paragraphs[0]
                paragraph.alignment = 1

def add_content_cogcaptable_remark(doc, cogcap_output):
    """Adds remarks to the cognitive capacity table."""
    if not isinstance(cogcap_output, str):
        logging.warning("cogcap_output is not a string.")
        return

    table = _safe_get_table(doc, COGCAP_TABLE_INDEX)
    if not table:
        return

    remark_cell = _safe_get_cell(table, 2, 1)  # Row 2 (third row)
    if not remark_cell:
        return

    _safe_set_text(remark_cell, cogcap_output)

def language_skills(doc, replacements_str, doc_obj): # Modified to accept doc object
    """Fills in language skills by replacing placeholders in paragraphs."""
    replacements = _safe_literal_eval(replacements_str, [])
    if not isinstance(replacements, list):
        logging.warning("Replacements is not a list.")
        return

    language_names = ["Dutch", "French", "English"]
    for index, language_name in enumerate(language_names):
        if index < len(replacements):
            proficiency_level = replacements[index]
            placeholder = f"{{prompt5_language_{language_name.lower()}}}" # Placeholder per language
            find_and_replace_placeholder(doc_obj, placeholder, proficiency_level) # Using find_and_replace_placeholder
        else:
            logging.warning(f"No proficiency level provided for {language_name}.")


def add_icons2(doc, list_scores):
    """Adds icons to the profile review tables (MCP version)."""
    if not isinstance(list_scores, list):
        logging.warning("list_scores is not a list.")
        return
    table_no_start = FIRST_ICONS_TABLE
    score_index = 0
    for table_no_offset in range(NUM_ICONS_TABLES): # Number of tables
      table_no = table_no_start + table_no_offset
      table = _safe_get_table(doc, table_no)
      if not table:
          continue  # Skip to next table

      for row_no in range(1, len(table.rows)): #Start from row 1
        if score_index < len(list_scores): # Check if scores remain
            cell = _safe_get_cell(table, row_no, 0) # Get the first cell
            if cell:
                add_icon_to_cell(cell, list_scores[score_index]) # Use function
                score_index += 1
        else:
            return

def add_icon_to_cell(cell, score):
    """Adds an icon based on the score to a cell (modified for MCP)."""
    if not isinstance(score, int):
        logging.warning(f"Invalid score type: {type(score)}. Expected int.")
        return

    if cell is None:
        logging.warning("add_icon_to_cell called with None cell.")
        return

    _safe_set_text(cell, "")  # Clear cell content

    run = cell.paragraphs[0].add_run()
    if score == -1:
        run.add_picture(resource_path("resources/improvement.png"), width=Inches(.3))
    elif score == 0:
        run.add_picture(resource_path("resources/average.png"), width=Inches(.3))
    elif score == 1:
        run.add_picture(resource_path("resources/strong.png"), width=Inches(.3))
    else:
        logging.warning(f"Invalid score value: {score}")

def conclusion(doc, column, list_items_str):
    """Adds conclusion bullet points, removing brackets and quotes."""
    table = _safe_get_table(doc, CONCLUSION_TABLE_INDEX)
    if not table:
        return

    list_items = _safe_literal_eval(list_items_str, [])
    if not isinstance(list_items, list):
        logging.warning("list_items could not be parsed as a list.")
        return

    cell = _safe_get_cell(table, 1, column)
    if not cell:
        return

    _safe_set_text(cell, "") # Clear cell content

    # --- Cleaning logic from the other script adapted here ---
    if isinstance(list_items, list):
        cleaned_items = []
        for item in list_items:
            if isinstance(item, str):
                cleaned_item = item.strip('"').strip("'") # Remove leading/trailing quotes
                cleaned_items.append(cleaned_item)
            else:
                cleaned_items.append(item) # Keep non-string items as they are
        list_points = cleaned_items
    else: # Fallback in case list_items is not a list after all (shouldn't happen)
        list_points = list_items


    for point in list_points:
        if isinstance(point, str):
            _safe_add_paragraph(cell, f'\t -{point}')

def replace_and_format_header_text(doc, new_text):
    """Replaces and formats header text."""
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            if '***' in paragraph.text:
                paragraph.text = paragraph.text.replace('***', new_text)
                for run in paragraph.runs:
                    run.font.name = 'Montserrat SemiBold'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(*(0xED, 0x6B, 0x55))
                    run.bold = True
                    run.italic = False
                    rFonts = OxmlElement('w:rFonts')
                    rFonts.set(qn('w:ascii'), 'Montserrat SemiBold')
                    rFonts.set(qn('w:hAnsi'), 'Montserrat SemiBold')
                    run._element.rPr.append(rFonts)

def replace_placeholder_docx(doc, placeholder, replacement, font_name='Montserrat', font_size=10): # corrected function name
    """Replaces placeholders with custom font."""
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            inline = paragraph.runs
            for i in range(len(inline)):
                if placeholder in inline[i].text:
                    inline[i].text = inline[i].text.replace(placeholder, replacement)
                    inline[i].font.name = font_name
                    inline[i].font.size = Pt(font_size)

def open_file(file_path):
    """Opens file based on OS (kept for consistency)."""
    if os.name == 'nt':
        os.startfile(file_path)
    elif os.name == 'posix':
        os.system(f'open "{file_path}"')
