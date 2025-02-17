import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import ast
import json
import re
import logging

# --- Configure Logging ---
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    filename="report_generation.log",
    filemode="a"
)

# --- Constants ---
DETAILS_TABLE_INDEX = 0
COGCAP_TABLE_INDEX = 1
CONCLUSION_TABLE_INDEX = 2
HUMAN_SKILLS_START_TABLE = 4
HUMAN_SKILLS_TABLE_COUNT = 5
TECH_SKILLS_START_TABLE = 9
TECH_SKILLS_TABLE_COUNT = 5
DATA_TOOLS_TABLE_START = 15
DATA_TOOLS_ITEMS_PER_TABLE = 5
INTERESTS_TABLE_INDEX = 16
LANGUAGE_SKILLS_TABLE_INDEX = 14

def resource_path(relative_path):
    """Get absolute path to resource, works for dev and PyInstaller."""
    try:
        base_path = sys._MEIPASS
    except AttributeError:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# --- Helper Functions ---
def _safe_get_table(doc, table_index, default=None):
    """Safely retrieves a table."""
    try: return doc.tables[table_index]
    except IndexError:
        logging.warning(f"Table {table_index} not found.")
        return default

def _safe_get_cell(table, row_index, col_index, default=None):
    """Safely retrieves a cell."""
    try: return table.cell(row_index, col_index)
    except IndexError:
        logging.warning(f"Cell ({row_index}, {col_index}) not found.")
        return default

def _safe_set_text(cell, text):
    """Safely sets cell text, clearing existing content."""
    if cell:
        for p in cell.paragraphs:
            p = p._element
            p.getparent().remove(p)
        paragraph = cell.add_paragraph()
        run = paragraph.add_run(str(text))
        run.font.name = 'Montserrat Light'
        run.font.size = Pt(10)

def _safe_add_paragraph(cell, text):
     """ Safely add paragraphs """
     if cell:
        paragraph = cell.add_paragraph(text)
        run = paragraph.runs[0]
        run.font.name = 'Montserrat Light'
        run.font.size = Pt(10)

        r = run._element
        rPr = r.rPr
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.append(rPr)

        rFonts = OxmlElement('w:rFonts')
        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', 'Montserrat Light')
        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi', 'Montserrat Light')
        rPr.append(rFonts)

def _safe_literal_eval(s, default=None):
    """Safely evaluates a string as a Python literal, removing backslashes."""
    try:
        # Remove backslashes before evaluation
        s = s.replace("\\", "")
        return ast.literal_eval(s)
    except (SyntaxError, ValueError):
        logging.error(f"Error evaluating: {s}")
        return default

def find_and_replace_placeholder(doc, placeholder, replacement_text, font_name='Montserrat Light', font_size=10):
    """Finds and replaces a placeholder in the document (paragraphs and tables)."""
    placeholder_found = False  # Flag

    # 1. Search Paragraphs
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            placeholder_found = True
            # Replace text in runs, preserving existing runs (and their formatting)
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(replacement_text))
                    run.font.name = font_name
                    run.font.size = Pt(font_size)

    # 2. Search Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    placeholder_found = True
                    _safe_set_text(cell, replacement_text) # Use safe set (clears and sets font)

    if not placeholder_found:
        logging.warning(f"Placeholder '{placeholder}' not found.")

# --- Main Functions ---
def clean(text):
    """Cleans input text."""
    return re.sub(r'[\【】`]|(```python)|(\*\*)', '', str(text)).strip() if isinstance(text, str) else text

def strip_extra_quotes(input_string):
    """Removes leading/trailing double quotes."""
    if isinstance(input_string, str) and input_string.startswith('"') and input_string.endswith('"'):
        return input_string[1:-1]
    return input_string

def clean_up(loc_dic):
    """Loads and cleans JSON data."""
    try:
        with open(loc_dic, 'r', encoding='utf-8') as f:
            loaded_data = json.load(f)

        # Remove backslashes from string values before further processing
        cleaned_data = {}
        for key, value in loaded_data.items():
            if isinstance(value, str):
                cleaned_data[key] = clean(value.replace("\\", ""))
            elif isinstance(value, list):  # Handle lists of strings
                 cleaned_data[key] = [clean(item.replace("\\","")) if isinstance(item,str) else item for item in value ]
            else:
                cleaned_data[key] = clean(value)  # Apply clean to other types as needed

        return cleaned_data
    except (FileNotFoundError, json.JSONDecodeError) as e:
        logging.error(f"Error loading/cleaning JSON: {e}")
        return {}

def update_document(output_dic, name, assessor, gender, program):
    """Updates the Word document."""
    try:
        doc = Document(resource_path('resources/Assessment_report_Data_chiefs.docx'))
    except Exception as e:
        logging.error(f"Failed to open template: {e}")
        return None

    # --- Static Content (still using dedicated functions) ---
    add_content_detailstable(doc, [name, "", program, "", ""])
    replace_and_format_header_text(doc, name)
    replace_placeholder_in_docx(doc, '***', name.split()[0], font_name='Montserrat Light')
    replace_placeholder_in_docx(doc, 'ASSESSOR', assessor.upper(), font_name='Montserrat Light')

    # --- Dynamic Content (using find_and_replace_placeholder) ---
    #  'prompt6a_conqual' and 'prompt6b_conimprov' are now handled by conclusion()
    dynamic_prompts = [
        'prompt2_firstimpr', 'prompt3_personality',
        'prompt4_cogcap_remarks'
    ]

    for prompt_key in dynamic_prompts:
        replacement = output_dic.get(prompt_key, "")
        if prompt_key in ['prompt2_firstimpr', 'prompt3_personality', 'prompt4_cogcap_remarks']:
            replacement = replacePiet(replacement, name, gender)
        find_and_replace_placeholder(doc, f"{{{prompt_key}}}", replacement)

    # --- Table/Specific Location Content ---
    add_content_cogcaptable(doc, output_dic.get('prompt4_cogcap_scores', "[]"))
    language_skills(doc, output_dic.get('prompt5_language', "[]"))

    # --- Conclusion Table ---
    conclusion(doc, 0, output_dic.get('prompt6a_conqual', "[]"))  # Strengths to column 0
    conclusion(doc, 1, output_dic.get('prompt6b_conimprov', "[]"))  # Improvements to column 1

        # --- Interests ---
    interests_str = output_dic.get('prompt9_interests', "")
    add_interests_table(doc, interests_str)

    # Profile review (icons)
    qual_scores_str = output_dic.get('prompt7_qualscore_data', "[]")
    qual_scores = _safe_literal_eval(qual_scores_str, [])
    if isinstance(qual_scores, list) and len(qual_scores) >= 23:
        add_icons_data_chief(doc, qual_scores[:18])
        add_icons_data_chief_2(doc, qual_scores[18:23])
    else:
        logging.warning("Invalid qual_scores data.")

    # Data tools (icons)
    data_tools_str = output_dic.get('prompt8_datatools', "[]")
    data_tools_scores = _safe_literal_eval(data_tools_str, [])
    if isinstance(data_tools_scores, list):
        add_icons_data_tools(doc, data_tools_scores)
    else:
        logging.warning("Invalid data_tools_scores data.")

    # --- Save Document ---
    current_time = datetime.now()
    formatted_time = current_time.strftime("%m%d%H%M")
    updated_doc_path = f"Assessment Report - {name} - {formatted_time}.docx"
    try:
        doc.save(updated_doc_path)
        return updated_doc_path
    except Exception as e:
        logging.error(f"Failed to save document: {e}")
        return None

def format_datatools_output(datatools_json_string):
    """Formats data tools output from JSON string."""
    try:
        datatools_dict = ast.literal_eval(datatools_json_string)
        formatted_text = ""
        for tool, level in datatools_dict.items():
            formatted_text += f"- {tool}: {level}\n"
        return formatted_text.strip()
    except (ValueError, SyntaxError):
        return "Could not parse data tools information."


def format_interests_output(interests_json_string):
    """Formats interests output from JSON string."""
    try:
        interests_list = ast.literal_eval(interests_json_string)
        formatted_text = ""
        for interest in interests_list:
            formatted_text += f"- {interest}\n"
        return formatted_text.strip()
    except (ValueError, SyntaxError):
        return "Could not parse interests information."

def replacePiet(text, name, gender):
    """Replaces 'Piet' and handles gender-specific pronouns."""
    if not isinstance(text, str):
        return ""  # Return empty string if not a string

    text = text.replace("Piet", name.split()[0])
    text = re.sub(r'\bthe trainee\b', name.split()[0], text, flags=re.IGNORECASE)

    if gender == 'M':
        replacements = {
            "She": "He", "she": "he", "Her": "Him", "her": "him",
            "Hers": "His", "hers": "his", "Herself": "Himself", "herself": "himself"
        }
    elif gender == 'F':
        replacements = {
            "He": "She", "he": "she", "Him": "Her", "him": "her",
            "His": "Her", "his": "her", "Himself": "Herself", "himself": "herself"
        }
    else:
        return text  # Return original text if gender is unknown

    for female, male in replacements.items():
        text = re.sub(r'\b' + re.escape(female) + r'\b', male, text)
    return text

def restructure_date(date_str):
    """Restructures date string to DD-MM-YYYY format."""
    date_str = date_str.replace('/', '-')

    try:
        datetime.strptime(date_str, '%d-%m-%Y')
        return date_str
    except ValueError:
        try:
            date_obj = datetime.strptime(date_str, '%Y-%m-%d')
            return date_obj.strftime('%d-%m-%Y')
        except ValueError:
            return ''


def add_content_detailstable(doc, personal_details):
    """Adds personal details to the first table."""
    table = _safe_get_table(doc, DETAILS_TABLE_INDEX)
    if not table:
        return

    if not isinstance(personal_details, list):
        logging.warning("personal_details is not a list.")
        return

    if len(personal_details) == 1 and all(isinstance(ele, str) for ele in personal_details):
        personal_details = personal_details[0].split(',')

    for row_index, row in enumerate(table.rows):  # Use enumerate
        if len(row.cells) > 1:
            first_cell_text = row.cells[0].text.strip()
            second_cell_text = row.cells[1].text.strip()

            if first_cell_text == "Name candidate" and second_cell_text == ":":
                cell = _safe_get_cell(table, row_index, 2)  # Use row_index
                _safe_set_text(cell, personal_details[0] if len(personal_details) > 0 else '')

            if first_cell_text == "Date of birth" and second_cell_text == ":":
                cell = _safe_get_cell(table, row_index, 2)  # Use row_index
                _safe_set_text(cell, restructure_date(personal_details[1]) if len(personal_details) > 1 else '')

            if first_cell_text == "Position" and second_cell_text == ":":
                cell = _safe_get_cell(table, row_index, 2)  # Use row_index
                _safe_set_text(cell, personal_details[2] if len(personal_details) > 2 else '')

            if first_cell_text == "Assessment date" and second_cell_text == ":":
                cell = _safe_get_cell(table, row_index, 2)  # Use row_index
                _safe_set_text(cell, restructure_date(personal_details[3]) if len(personal_details) > 3 else '')

            if first_cell_text == "Pool" and second_cell_text == ":":
                cell = _safe_get_cell(table, row_index, 2)  # Use row_index
                _safe_set_text(cell, personal_details[4] if len(personal_details) > 4 else '')

def add_content_cogcaptable(doc, scores_str):
    """Adds cognitive capacity scores."""
    table = _safe_get_table(doc, COGCAP_TABLE_INDEX)
    if not table:
        return

    scores = _safe_literal_eval(scores_str, [])
    if not isinstance(scores, list) or len(scores) != 6:
        logging.warning("Invalid scores data. Expected a list of 6 numbers.")
        return

    for i in range(6):
        cell = _safe_get_cell(table, 1, i + 1)  # Row 1 (second row)
        if cell:
            if i == 0:
                _safe_set_text(cell, scores[i])# Already handles string conversion
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.bold = True
                run.underline = True
                paragraph.alignment = 1
            else:
                _safe_set_text(cell, scores[i])
                paragraph = cell.paragraphs[0]
                paragraph.alignment = 1



def add_content_cogcaptable_remark(doc, cogcap_output):
    """Adds remarks to the cognitive capacity table."""
    if not isinstance(cogcap_output, str):
        logging.warning("cogcap_output is not a string.")
        return

    table = _safe_get_table(doc, COGCAP_TABLE_INDEX)
    if not table:
        return

    remark_cell = _safe_get_cell(table, 2, 1)  # Row 2 (third row)
    if not remark_cell:
        return

    _safe_set_text(remark_cell, cogcap_output)


def language_skills(doc, replacements_str):
    """Fills in language skills levels in the Language Skills table."""
    replacements = _safe_literal_eval(replacements_str, [])
    if not isinstance(replacements, list):
        logging.warning("Replacements is not a list.")
        return

    table = _safe_get_table(doc, LANGUAGE_SKILLS_TABLE_INDEX)
    if not table:
        return

    language_names = ["Dutch", "French", "English"]
    for index, language_name in enumerate(language_names):
        if index < len(replacements):
            proficiency_level = replacements[index]
            cell = _safe_get_cell(table, index + 2, 0) # Row index +2 (skip header)
            _safe_set_text(cell, proficiency_level)

        else:
            logging.warning(f"No proficiency level provided for {language_name}.")


def add_icons_data_chief(doc, list_scores):
    """Adds icons to Human Skills tables."""
    if not isinstance(list_scores, list):
        logging.warning("list_scores is not a list.")
        return

    score_index = 0
    for table_no_offset in range(HUMAN_SKILLS_TABLE_COUNT):
        table_no = HUMAN_SKILLS_START_TABLE + table_no_offset
        table = _safe_get_table(doc, table_no)
        if not table:
            continue  # Skip to the next table if this one is not found

        for row_no in range(1, len(table.rows)):
            if score_index < len(list_scores):
                cell = _safe_get_cell(table, row_no, 0)
                if cell and cell.text.strip().startswith("AA"):
                    add_icon_to_cell(cell, list_scores[score_index])
                    score_index += 1
            else:
                break  # All scores processed

def add_icons_data_chief_2(doc, list_scores):
    """Adds icons to Technical Skills tables."""
    if not isinstance(list_scores, list):
        logging.warning("list_scores is not a list.")
        return

    score_index = 0
    for table_no_offset in range(TECH_SKILLS_TABLE_COUNT):
        table_no = TECH_SKILLS_START_TABLE + table_no_offset
        table = _safe_get_table(doc, table_no)
        if not table:
            continue # Skip

        for row_no in range(1, len(table.rows)):
            if score_index < len(list_scores):
                cell = _safe_get_cell(table, row_no, 0)
                if cell and cell.text.strip().startswith("AA"):
                    add_icon_to_cell(cell, list_scores[score_index])
                    score_index += 1
            else:
                break # All scores

def add_icons_data_tools(doc, list_scores):
    """Adds icons to Data Tools tables."""
    if not isinstance(list_scores, list):
        logging.warning("list_scores is not a list.")
        return

    for i in range(len(list_scores)):
        table_no = DATA_TOOLS_TABLE_START + (i // DATA_TOOLS_ITEMS_PER_TABLE)
        row_no = (i % DATA_TOOLS_ITEMS_PER_TABLE) + 2 # +2 to skip header

        table = _safe_get_table(doc, table_no)
        if not table:
            continue

        cell = _safe_get_cell(table, row_no, 0)
        if cell:
            add_icon_to_cell(cell, list_scores[i])


def add_icon_to_cell(cell, score):
    """Adds an icon based on the score to a cell."""
    if not isinstance(score, int):
        logging.warning(f"Invalid score type: {type(score)}. Expected int.")
        return

    if cell is None: # Check
        logging.warning("add_icon_to_cell called with None cell.")
        return
    # Clear existing content using _safe_set_text
    _safe_set_text(cell, "")

    run = cell.paragraphs[0].add_run()  # Use existing, cleared paragraph
    if score == -1:
        run.add_picture(resource_path("resources/improvement.png"), width=Inches(.3))
    elif score == 0:
        run.add_picture(resource_path("resources/average.png"), width=Inches(.3))
    elif score == 1:
        run.add_picture(resource_path("resources/strong.png"), width=Inches(.3))
    else:
        logging.warning(f"Invalid score value: {score}")

def add_interests_table(doc, interests_text):
    """Fills in interests into the Interests Table as comma-separated text, handling strings directly."""
    table = _safe_get_table(doc, INTERESTS_TABLE_INDEX)
    if not table:
        return

    # --- Key Changes Here ---
    if isinstance(interests_text, str):
        # Remove brackets and any extra whitespace, then split by comma
        interests_list = [s.strip() for s in interests_text.strip("[]").split(",") if s.strip()]
        # Remove quotes around each interest
        interests_list = [s.strip('"').strip("'") for s in interests_list]
        interests_string = ', '.join(interests_list)
    else:
        logging.warning("interests_text is not a string.")
        interests_string = ""


    # Get the first cell of the interests table
    cell = _safe_get_cell(table, 1, 0)  # Assuming interests are in the second row, first column
    if cell:
        _safe_set_text(cell, interests_string)
    else:
        logging.warning("Could not find cell to add interests text.")


def conclusion(doc, column, list_items):
    """Adds conclusion bullet points to the specified column, handling lists directly."""
    table = _safe_get_table(doc, CONCLUSION_TABLE_INDEX)
    if not table:
        return

    # Directly use the output of _safe_literal_eval
    if isinstance(list_items, str):  # If it's still a string, try to evaluate
        list_items = _safe_literal_eval(list_items, [])

    if not isinstance(list_items, list):
        logging.warning("list_items could not be parsed as a list")
        return

    cell = _safe_get_cell(table, 1, column)
    if not cell:
        return
    _safe_set_text(cell, "")  # Clear existing content

    for point in list_items:
        if isinstance(point, str):
            _safe_add_paragraph(cell, f'\t -{point}')  # Use safe function


# Last style improvements
def replace_and_format_header_text(doc, new_text):
    """Replaces header text and formats it."""
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            if '***' in paragraph.text:
                paragraph.text = paragraph.text.replace('***', new_text)
                for run in paragraph.runs:
                    run.font.name = 'Montserrat SemiBold'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(*(0xED, 0x6B, 0x55))
                    run.bold = True
                    run.italic = False
                    rFonts = OxmlElement('w:rFonts')
                    rFonts.set(qn('w:ascii'), 'Montserrat SemiBold')
                    rFonts.set(qn('w:hAnsi'), 'Montserrat SemiBold')
                    run._element.rPr.append(rFonts)


def replace_placeholder_in_docx(doc, placeholder, replacement, font_name='Montserrat', font_size=10):
    """Replaces a placeholder in the document with custom font."""
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            inline = paragraph.runs
            for i in range(len(inline)):
                if placeholder in inline[i].text:
                    inline[i].text = inline[i].text.replace(placeholder, replacement)
                    inline[i].font.name = font_name
                    inline[i].font.size = Pt(font_size)

def open_file(file_path):  # No changes needed, kept for completeness
    """Opens file based on OS."""
    if os.name == 'nt':
        os.startfile(file_path)
    elif os.name == 'posix':
        os.system(f'open "{file_path}"')
